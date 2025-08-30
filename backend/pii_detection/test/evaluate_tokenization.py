
"""
分词评估脚本（测试目录）
"""

import os
import re
import sys
import logging
from typing import List, Tuple, Dict

# 将项目根目录加入 Python 路径，便于容器内导入
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, '..', '..'))  # /app
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

# Django 环境
import django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'compliance_agent.settings')
django.setup()

# 导入视图函数
from pii_detection.views import PiiDetectView

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TokenizationEvaluator:
    def __init__(self):
        self.pii_detector = PiiDetectView()
        
    def load_test_samples(self, test_file: str, clean_file: str) -> List[Tuple[str, str]]:
        import docx
        # 原始样本（docx）
        doc = docx.Document(test_file)
        test_text = '\n'.join([p.text for p in doc.paragraphs])
        # 人工分词（docx）
        doc2 = docx.Document(clean_file)
        clean_text = '\n'.join([p.text for p in doc2.paragraphs])
        
        separator = "=" * 50
        test_samples = [s.strip() for s in test_text.split(separator) if s.strip()]
        clean_samples = [s.strip() for s in clean_text.split(separator) if s.strip()]
        
        if len(test_samples) != len(clean_samples):
            logger.warning(f"样本数量不匹配: 测试样本{len(test_samples)}个, 人工分词{len(clean_samples)}个")
        return list(zip(test_samples, clean_samples))

    def extract_tokens(self, text: str) -> List[str]:
        return [t for t in text.split() if t]

    def get_auto_tokens(self, text: str) -> List[str]:
        try:
            structured_text = self.pii_detector.process_text_with_standardization(text)
            tokens: List[str] = []
            for sentence_tokens in structured_text.get('tokenized_sentences', []):
                if sentence_tokens.startswith('[') and sentence_tokens.endswith(']'):
                    tokens_str = sentence_tokens[1:-1]
                    tokens.extend([t.strip().strip("'") for t in tokens_str.split("', '")])
            return tokens
        except Exception as e:
            logger.error(f"自动分词失败: {e}")
            return text.split()

    def _tokens_to_plain_and_boundaries(self, tokens: List[str]) -> Tuple[str, List[int]]:
        """
        将 tokens 转为无空格的串，以及每个词末尾的边界位置集合（不含句末）。
        边界位置按字符计数，范围 [1, len(plain)-1]。
        """
        plain = ''.join(tokens)
        boundaries = []
        acc = 0
        for tok in tokens[:-1]:
            acc += len(tok)
            boundaries.append(acc)
        return plain, boundaries

    def boundary_metrics(self, raw_text: str, manual_tokens: List[str], auto_tokens: List[str]) -> Tuple[Dict[str, float], bool]:
        """
        基于切分边界的指标（Precision/Recall/F1/Accuracy）。
        - 若 manual 与 auto 在去空格后的底层串不一致，则返回 fallback=False，提示外层使用集合对比退化方案。
        """
        manual_plain, m_bounds = self._tokens_to_plain_and_boundaries(manual_tokens)
        auto_plain, a_bounds = self._tokens_to_plain_and_boundaries(auto_tokens)
        if manual_plain != auto_plain:
            # 底层串不一致，无法公平比较边界
            return ({'accuracy': 0.0, 'precision': 0.0, 'recall': 0.0, 'f1': 0.0}, False)
        L = len(manual_plain)
        total_positions = max(L - 1, 0)
        m_set = set(m_bounds)
        a_set = set(a_bounds)
        tp = len(m_set & a_set)
        fp = len(a_set - m_set)
        fn = len(m_set - a_set)
        tn = total_positions - tp - fp - fn
        accuracy = (tp + tn) / total_positions if total_positions > 0 else 0.0
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0.0
        return ({'tp': tp, 'fp': fp, 'fn': fn, 'tn': tn, 'accuracy': accuracy, 'precision': precision, 'recall': recall, 'f1': f1}, True)

    # 基于 token 集合的粗略评估
    def set_based_metrics(self, auto_tokens: List[str], manual_tokens: List[str]) -> Dict[str, float]:
        auto_set = set(auto_tokens)
        manual_set = set(manual_tokens)
        tp = len(auto_set & manual_set)
        fp = len(auto_set - manual_set)
        fn = len(manual_set - auto_set)
        tn = 0
        total = tp + tn + fp + fn
        accuracy = (tp + tn) / total if total else 0
        precision = tp / (tp + fp) if (tp + fp) else 0
        recall = tp / (tp + fn) if (tp + fn) else 0
        f1 = 2 * (precision * recall) / (precision + recall) if (precision + recall) else 0
        iou = tp / (tp + fp + fn) if (tp + fp + fn) else 0  # Jaccard
        return {'tp': tp, 'fp': fp, 'fn': fn, 'tn': tn, 'accuracy': accuracy, 'precision': precision, 'recall': recall, 'f1': f1, 'iou': iou}

    def evaluate_all_samples(self, test_file: str, clean_file: str) -> Dict:
        samples = self.load_test_samples(test_file, clean_file)
        total_boundary = {'tp': 0, 'fp': 0, 'fn': 0, 'tn': 0}
        total_set = {'tp': 0, 'fp': 0, 'fn': 0, 'tn': 0}
        sample_results = []
        boundary_applicable = 0
        for i, (raw, manual) in enumerate(samples):
            manual_tokens = self.extract_tokens(manual)
            auto_tokens = self.get_auto_tokens(raw)

            # 边界评估（优先）
            b_metrics, ok = self.boundary_metrics(raw, manual_tokens, auto_tokens)
            if ok:
                boundary_applicable += 1
                for k in ('tp', 'fp', 'fn', 'tn'):
                    total_boundary[k] += b_metrics[k]

            # 集合评估（保留）
            s_metrics = self.set_based_metrics(auto_tokens, manual_tokens)
            for k in ('tp', 'fp', 'fn', 'tn'):
                total_set[k] += s_metrics[k]

            sample_results.append({
                'sample_id': i + 1,
                'manual_tokens': manual_tokens,
                'auto_tokens': auto_tokens,
                'metrics_boundary': b_metrics if ok else None,
                'metrics_set': s_metrics,
                'boundary_used': ok
            })

        # 汇总：边界
        summary = {'boundary': None, 'set_based': None, 'boundary_applicable': boundary_applicable}
        if boundary_applicable > 0:
            tp, fp, fn, tn = total_boundary['tp'], total_boundary['fp'], total_boundary['fn'], total_boundary['tn']
            total_positions = tp + fp + fn + tn
            accuracy = (tp + tn) / total_positions if total_positions else 0.0
            precision = tp / (tp + fp) if (tp + fp) else 0.0
            recall = tp / (tp + fn) if (tp + fn) else 0.0
            f1 = 2 * precision * recall / (precision + recall) if (precision + recall) else 0.0
            summary['boundary'] = {'tp': tp, 'fp': fp, 'fn': fn, 'tn': tn, 'accuracy': accuracy, 'precision': precision, 'recall': recall, 'f1': f1}

        # 汇总：集合
        tp, fp, fn, tn = total_set['tp'], total_set['fp'], total_set['fn'], total_set['tn']
        total = tp + fp + fn + tn
        accuracy = (tp + tn) / total if total else 0
        precision = tp / (tp + fp) if (tp + fp) else 0
        recall = tp / (tp + fn) if (tp + fn) else 0
        f1 = 2 * (precision * recall) / (precision + recall) if (precision + recall) else 0
        iou = tp / (tp + fp + fn) if (tp + fp + fn) else 0
        summary['set_based'] = {'tp': tp, 'fp': fp, 'fn': fn, 'tn': tn, 'accuracy': accuracy, 'precision': precision, 'recall': recall, 'f1': f1, 'iou': iou}

        return {'summary': summary, 'sample_results': sample_results, 'sample_count': len(samples)}

    def print_results(self, results: Dict):
        s = results['summary']
        print("=" * 60)
        print("分词评估结果（常用：切分边界法）")
        print("=" * 60)
        print(f"样本数量: {results['sample_count']}  （可进行边界评估样本: {s['boundary_applicable']}）")
        print()
        if s['boundary']:
            m = s['boundary']
            print("切分边界总体:")
            print(f"Accuracy={m['accuracy']:.4f}, Precision={m['precision']:.4f}, Recall={m['recall']:.4f}, F1={m['f1']:.4f}")
        else:
            print("切分边界总体: 无（底层串不一致，无法计算）")
        print()
        m2 = s['set_based']
        print("集合对比（直接对比分词结果）:")
        print(f"Precision={m2['precision']:.4f}, Recall={m2['recall']:.4f}, F1={m2['f1']:.4f}, IoU={m2['iou']:.4f}")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    test_file = os.path.join(base_dir, 'test1.docx')
    clean_file = os.path.join(base_dir, 'test1_clean.docx')
    evaluator = TokenizationEvaluator()
    results = evaluator.evaluate_all_samples(test_file, clean_file)
    evaluator.print_results(results)
    out_file = os.path.join(base_dir, 'tokenization_evaluation_results.txt')
    with open(out_file, 'w', encoding='utf-8') as f:
        f.write("分词评估详细结果\n")
        f.write("=" * 60 + "\n\n")
        s = results['summary']
        f.write(f"样本数量: {results['sample_count']}\n")
        f.write(f"可进行边界评估样本: {s['boundary_applicable']}\n\n")
        if s['boundary']:
            m = s['boundary']
            f.write("切分边界总体:\n")
            f.write(f"Accuracy={m['accuracy']:.4f}, Precision={m['precision']:.4f}, Recall={m['recall']:.4f}, F1={m['f1']:.4f}\n\n")
        else:
            f.write("切分边界总体: 无（底层串不一致，无法计算）\n\n")
        m2 = s['set_based']
        f.write("集合对比（参考）:\n")
        f.write(f"Precision={m2['precision']:.4f}, Recall={m2['recall']:.4f}, F1={m2['f1']:.4f}, IoU={m2['iou']:.4f}\n\n")
        f.write("样本详细结果:\n")
        f.write("-" * 60 + "\n")
        for sres in results['sample_results']:
            f.write(f"样本 {sres['sample_id']}:\n")
            if sres['boundary_used'] and sres['metrics_boundary']:
                mb = sres['metrics_boundary']
                f.write(f"  边界: Acc={mb.get('accuracy',0):.4f}, P={mb.get('precision',0):.4f}, R={mb.get('recall',0):.4f}, F1={mb.get('f1',0):.4f}\n")
            else:
                f.write("  边界: 无\n")
            ms = sres['metrics_set']
            f.write(f"  集合: P={ms['precision']:.4f}, R={ms['recall']:.4f}, F1={ms['f1']:.4f}, IoU={ms['iou']:.4f}\n")
            f.write("-" * 60 + "\n")
    print(f"详细结果已保存到: {out_file}")

if __name__ == '__main__':
    main()
